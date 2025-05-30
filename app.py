import os
from flask import Flask, render_template, request, send_file, jsonify, url_for, make_response
import pandas as pd
from werkzeug.utils import secure_filename
import zipfile
from docx import Document
import tempfile
from docx2pdf import convert
import pythoncom
import shutil
import time
import io
import uuid
from datetime import datetime, timedelta

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['GENERATED_FILES'] = 'generated_files'  # New folder for generated files
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['GENERATED_FILES'], exist_ok=True)

# Store generated files with their IDs and creation time
generated_files = {}

def cleanup_old_files():
    """Clean up files older than 1 hour"""
    current_time = datetime.now()
    files_to_remove = []
    
    for file_id, file_info in generated_files.items():
        if current_time - file_info['created_at'] > timedelta(hours=1):
            try:
                os.remove(file_info['path'])
                files_to_remove.append(file_id)
            except:
                pass
    
    for file_id in files_to_remove:
        del generated_files[file_id]

def extract_placeholders(doc):
    import re
    placeholders = set()
    regex = re.compile(r'\{[^{}]+\}')

    def process_element(element):
        text = ''.join(run.text for run in element.runs)
        matches = regex.findall(text)
        for m in matches:
            placeholders.add(m)

    for paragraph in doc.paragraphs:
        process_element(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_element(paragraph)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_element(paragraph)
        for paragraph in section.footer.paragraphs:
            process_element(paragraph)

    return sorted(list(placeholders))

def replace_placeholders(doc, data):
    def _process_paragraph(paragraph, data):
        runs = paragraph.runs
        text = ''.join([run.text for run in runs])

        for ph, val in data.items():
            if ph not in text:
                continue

            start_idx = 0
            while True:
                start_idx = text.find(ph, start_idx)
                if start_idx == -1:
                    break

                current_pos = 0
                start_run = None
                end_run = None

                for i, run in enumerate(runs):
                    run_end = current_pos + len(run.text)
                    if current_pos <= start_idx < run_end:
                        start_run = i
                    if current_pos <= (start_idx + len(ph)) <= run_end:
                        end_run = i
                        break
                    current_pos = run_end

                if start_run is not None and end_run is not None:
                    merged_text = ''.join(run.text for run in runs[start_run:end_run + 1])
                    replaced_text = merged_text.replace(ph, str(val))
                    runs[start_run].text = replaced_text
                    for run in runs[start_run + 1:end_run + 1]:
                        run.text = ""

                start_idx += len(ph)

    for paragraph in doc.paragraphs:
        _process_paragraph(paragraph, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _process_paragraph(paragraph, data)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _process_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            _process_paragraph(paragraph, data)

def get_column_names(file_path):
    try:
        df = pd.read_excel(file_path)
        return df.columns.tolist()
    except Exception as e:
        return []

def get_column_values(file_path, column):
    try:
        df = pd.read_excel(file_path)
        if column not in df.columns:
            return []
        # Get unique values and convert to list, handling NaN values
        values = df[column].dropna().unique().tolist()
        # Convert all values to strings for consistency
        return [str(value) for value in values]
    except Exception as e:
        return []

def compare_files(file1_path, file2_path, col1, col2):
    df1 = pd.read_excel(file1_path)
    df2 = pd.read_excel(file2_path)
    
    # Find records in file1 not in file2
    not_in_file2 = df1[~df1[col1].isin(df2[col2])]
    
    # Find records in file2 not in file1
    not_in_file1 = df2[~df2[col2].isin(df1[col1])]
    
    # Create output Excel file
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'comparison_result.xlsx')
    with pd.ExcelWriter(output_path) as writer:
        not_in_file2.to_excel(writer, sheet_name='Not in File 2', index=False)
        not_in_file1.to_excel(writer, sheet_name='Not in File 1', index=False)
    
    return output_path

def join_files(file1_path, file2_path, col1, col2):
    df1 = pd.read_excel(file1_path)
    df2 = pd.read_excel(file2_path)
    
    # Rename the join column in df2 to match df1 to avoid duplicates
    df2 = df2.rename(columns={col2: col1})
    
    # Perform the join operation (similar to SQL INNER JOIN)
    joined_df = pd.merge(df1, df2, on=col1, how='inner')
    
    # Create output Excel file
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'joined_result.xlsx')
    joined_df.to_excel(output_path, index=False)
    
    return output_path

def split_file(file_path, column, value):
    df = pd.read_excel(file_path)
    
    # Filter rows where the column value matches the selected value
    # Convert both to strings for comparison to handle different data types
    filtered_df = df[df[column].astype(str) == str(value)]
    
    # Create output Excel file
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], f'split_{value}.xlsx')
    filtered_df.to_excel(output_path, index=False)
    
    return output_path

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/compare')
def compare_page():
    return render_template('compare.html')

@app.route('/join')
def join_page():
    return render_template('join.html')

@app.route('/filter')
def filter_page():
    return render_template('filter.html')

@app.route('/split')
def split_by_column_page():
    return render_template('split.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    if 'excel' not in request.files or 'word' not in request.files:
        return jsonify({'error': 'Missing files'}), 400
    
    excel_file = request.files['excel']
    word_file = request.files['word']
    
    if excel_file.filename == '' or word_file.filename == '':
        return jsonify({'error': 'No files selected'}), 400

    # Save files temporarily
    excel_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(excel_file.filename))
    word_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(word_file.filename))
    
    excel_file.save(excel_path)
    word_file.save(word_path)

    try:
        # Read Excel columns
        df = pd.read_excel(excel_path)
        columns = df.columns.tolist()

        # Extract placeholders from Word
        doc = Document(word_path)
        placeholders = extract_placeholders(doc)

        return jsonify({
            'columns': columns,
            'placeholders': placeholders,
            'excel_path': excel_path,
            'word_path': word_path
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate', methods=['POST'])
def generate_documents():
    data = request.json
    excel_path = data.get('excel_path')
    word_path = data.get('word_path')
    mappings = data.get('mappings')
    key_column = data.get('key_column')
    export_formats = data.get('export_formats', {'docx': True, 'pdf': False})

    if not all([excel_path, word_path, mappings, key_column]):
        return jsonify({'error': 'Missing required data'}), 400

    if not any(export_formats.values()):
        return jsonify({'error': 'Please select at least one export format'}), 400

    try:
        df = pd.read_excel(excel_path)
        
        # Create a temporary directory for processing
        temp_dir = tempfile.mkdtemp(dir=app.config['UPLOAD_FOLDER'])
        try:
            # Process each row
            output_files = []
            filename_counters = {}  # Keep track of filename occurrences
            
            for index, row in df.iterrows():
                doc = Document(word_path)
                
                # Prepare data for replacement
                replace_data = {ph: str(row[col]) for ph, col in mappings.items()}
                replace_placeholders(doc, replace_data)
                
                base_filename = f"{row[key_column]}"
                
                # Handle duplicate filenames
                if base_filename in filename_counters:
                    filename_counters[base_filename] += 1
                    base_filename = f"{base_filename}_{filename_counters[base_filename]}"
                else:
                    filename_counters[base_filename] = 0
                
                # Save as DOCX if requested
                if export_formats.get('docx'):
                    docx_path = os.path.join(temp_dir, f"{base_filename}.docx")
                    doc.save(docx_path)
                    output_files.append(docx_path)
                
                # Convert to PDF if requested
                if export_formats.get('pdf'):
                    # Initialize COM for each conversion
                    pythoncom.CoInitialize()
                    try:
                        if not export_formats.get('docx'):
                            # If we don't want DOCX, save it temporarily
                            docx_path = os.path.join(temp_dir, f"{base_filename}_temp.docx")
                            doc.save(docx_path)
                        
                        pdf_path = os.path.join(temp_dir, f"{base_filename}.pdf")
                        convert(docx_path, pdf_path)
                        output_files.append(pdf_path)
                        
                        # Clean up temporary DOCX if we didn't want it
                        if not export_formats.get('docx'):
                            try:
                                os.remove(docx_path)
                            except OSError:
                                pass
                    finally:
                        pythoncom.CoUninitialize()

            # Create a temporary file for the ZIP
            zip_buffer = io.BytesIO()
            
            # Create the ZIP file in memory
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file in output_files:
                    if os.path.exists(file):  # Only add files that exist
                        zipf.write(file, os.path.basename(file))

            # Clean up the temporary processing directory
            try:
                shutil.rmtree(temp_dir)
            except OSError:
                pass  # Ignore errors during cleanup

            # Prepare the ZIP file for download
            zip_buffer.seek(0)
            
            # Create response with appropriate headers for automatic download
            response = make_response(zip_buffer.getvalue())
            response.headers['Content-Type'] = 'application/zip'
            response.headers['Content-Disposition'] = 'attachment; filename=generated_documents.zip'
            
            return response

        except Exception as e:
            # Clean up on error
            try:
                shutil.rmtree(temp_dir)
            except OSError:
                pass  # Ignore errors during cleanup
            raise e

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/column-values', methods=['POST'])
def column_values():
    if 'file' not in request.files:
        return jsonify({'error': 'Δεν έχει μεταφορτωθεί αρχείο'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Δεν έχει επιλεγεί αρχείο'}), 400
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': 'Παρακαλώ μεταφορτώστε αρχείο Excel'}), 400
    
    column = request.form.get('column')
    if not column:
        return jsonify({'error': 'Δεν έχει επιλεγεί στήλη'}), 400
    
    try:
        # Save the file temporarily
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Get unique values for the column
        values = get_column_values(file_path, column)
        
        # Clean up the temporary file
        os.remove(file_path)
        
        return jsonify({'values': values})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/compare', methods=['POST'])
def compare():
    data = request.json
    file1_path = os.path.join(app.config['UPLOAD_FOLDER'], data['file1'])
    file2_path = os.path.join(app.config['UPLOAD_FOLDER'], data['file2'])
    
    try:
        output_path = compare_files(file1_path, file2_path, data['col1'], data['col2'])
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/join', methods=['POST'])
def join():
    data = request.json
    file1_path = os.path.join(app.config['UPLOAD_FOLDER'], data['file1'])
    file2_path = os.path.join(app.config['UPLOAD_FOLDER'], data['file2'])
    
    try:
        output_path = join_files(file1_path, file2_path, data['col1'], data['col2'])
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/filter', methods=['POST'])
def filter_file():
    if 'file' not in request.files:
        return render_template('filter.html', error='Δεν έχει μεταφορτωθεί αρχείο')
    
    file = request.files['file']
    if file.filename == '':
        return render_template('filter.html', error='Δεν έχει επιλεγεί αρχείο')
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return render_template('filter.html', error='Παρακαλώ μεταφορτώστε αρχείο Excel')
    
    column = request.form.get('column')
    value = request.form.get('value')
    
    if not column or not value:
        return render_template('filter.html', error='Παρακαλώ επιλέξτε στήλη και εισάγετε τιμή')
    
    try:
        # Read the Excel file
        df = pd.read_excel(file)
        
        if column not in df.columns:
            return render_template('filter.html', error=f'Η στήλη "{column}" δεν βρέθηκε στο αρχείο')
        
        # Filter the dataframe
        filtered_df = df[df[column].astype(str) == str(value)]
        
        if filtered_df.empty:
            return render_template('filter.html', error=f'Δεν βρέθηκαν γραμμές όπου {column} ισούται με "{value}"')
        
        # Create output filename
        output_filename = f'filtered_{secure_filename(file.filename)}'
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        # Save the filtered dataframe
        filtered_df.to_excel(output_path, index=False)
        
        return render_template('filter.html', 
                             success=f'Βρέθηκαν {len(filtered_df)} γραμμές στις οποίες {column} ισούται με "{value}"',
                             download_link={'url': url_for('download_file', filename=output_filename),
                                          'filename': output_filename})
        
    except Exception as e:
        return render_template('filter.html', error=f'Σφάλμα κατά την επεξεργασία του αρχείου: {str(e)}')


@app.route('/get-columns', methods=['POST'])
def get_columns():
    if 'file' not in request.files:
        return jsonify({'error': 'Δεν έχει μεταφορτωθεί αρχείο'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Δεν έχει επιλεγεί αρχείο'}), 400
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': 'Παρακαλώ μεταφορτώστε αρχείο Excel'}), 400
    
    try:
        # Read the Excel file
        df = pd.read_excel(file)
        
        # Get column names
        columns = df.columns.tolist()
        
        return jsonify({'columns': columns})
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download/<path:filename>')
def download_file(filename):
    try:
        return send_file(
            os.path.join(app.config['UPLOAD_FOLDER'], filename),
            as_attachment=True
        )
    except Exception as e:
        return str(e), 404

@app.route('/split', methods=['POST'])
def split():
    if 'file' not in request.files:
        return render_template('split.html', error='Δεν έχει μεταφορτωθεί αρχείο')
    
    file = request.files['file']
    if file.filename == '':
        return render_template('split.html', error='Δεν έχει επιλεγεί αρχείο')
    
    if not file.filename.endswith(('.xlsx', '.xls')):
        return render_template('split.html', error='Παρακαλώ μεταφορτώστε αρχείο Excel')
    
    column = request.form.get('column')
    if not column:
        return render_template('split.html', error='Παρακαλώ επιλέξτε στήλη')
    
    try:
        # Read the Excel file
        df = pd.read_excel(file)
        
        if column not in df.columns:
            return render_template('split.html', error=f'Η στήλη "{column}" δεν βρέθηκε στο αρχείο')
        
        # Create a directory for split files
        output_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'split_files')
        os.makedirs(output_dir, exist_ok=True)
        
        # Split the dataframe by unique values in the selected column
        unique_values = df[column].unique()
        split_files = []
        
        for value in unique_values:
            # Create a clean filename from the value
            clean_value = str(value).replace('/', '_').replace('\\', '_')
            filename = f'split_{clean_value}.xlsx'
            filepath = os.path.join(output_dir, filename)
            
            # Save the filtered dataframe
            df[df[column] == value].to_excel(filepath, index=False)
            split_files.append(filepath)
        
        # Create a zip file containing all split files
        zip_filename = f'split_files_{secure_filename(file.filename)}.zip'
        zip_path = os.path.join(app.config['UPLOAD_FOLDER'], zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file_path in split_files:
                arcname = os.path.basename(file_path)
                zipf.write(file_path, arcname)
        
        # Clean up individual split files
        for file_path in split_files:
            os.remove(file_path)
        
        return render_template('split.html', 
                             success=f'Το αρχείο διαχωρίστηκε σε {len(split_files)} αρχεία με βάση τη στήλη "{column}"',
                             download_link={'url': url_for('download_file', filename=zip_filename),
                                          'filename': zip_filename})
        
    except Exception as e:
        return render_template('split.html', error=f'Σφάλμα κατά την επεξεργασία του αρχείου: {str(e)}')

@app.route('/excel-to-word')
def excel_to_word_page():
    return render_template('excel-to-word.html')

if __name__ == '__main__':
    # Get port from environment variable or default to 8000
    port = int(os.environ.get('PORT', 8000))
    # Run the app on all interfaces
    app.run(host='0.0.0.0', port=port) 